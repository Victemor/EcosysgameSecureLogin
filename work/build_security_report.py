from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\xampp\htdocs\secure-login")
OUT_DIR = ROOT / "output" / "documents"
QA_DIR = ROOT / "work" / "report_qa"
OUT_DIR.mkdir(parents=True, exist_ok=True)
QA_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "Informe_Pruebas_Seguridad_API_Ecosysgame_APA7.docx"

SCREENSHOTS = {
    "register": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-95ee20da-d791-4e27-96cf-162787f4333a.png"),
    "login": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-5fe58aab-dee0-4ea9-9568-dd9cd60c808a.png"),
    "profile": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-6877c4f1-7ceb-4f02-8fbe-95bc8193df66.png"),
    "no_token": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-c2bca68e-e343-414f-baad-32be1038873a.png"),
    "logout": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-d38577ba-bb41-42fc-8651-5fb1c0ffc7c4.png"),
    "revoked": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-1190815c-cdb0-4751-b0ce-fe88ef63c7b4.png"),
    "duplicate": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-8080cfab-0676-4a0a-8d62-a6b6e4c01bbc.png"),
    "weak": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-ec4b33ae-1056-4201-94e2-70efbcc7d77e.png"),
    "bad_login": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-84247a2d-9c18-4dee-b7c7-1c43536f4423.png"),
    "mass_assignment": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-3bf1a4f0-0e0e-478f-af0f-a6f2a40304be.png"),
    "rate_limit": Path(r"C:\Users\HEWLET~1\AppData\Local\Temp\codex-clipboard-e6c770c9-ac37-4270-a5f0-bd8bf13f1dc7.png"),
}


def font(run, size=12, bold=False, italic=False, name="Times New Roman", color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    font(run, size=12)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_para(doc, text="", bold_prefix=None, italic=False, align=None, first_line=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), bold=True)
        font(p.add_run(text[len(bold_prefix):]), italic=italic)
    else:
        font(p.add_run(text), italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    font(p.add_run(text))
    return p


def add_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], "D9EAD3")
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(header), size=font_size, bold=True, name="Arial")
        hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            font(p.add_run(str(value)), size=font_size, name="Arial")
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_figure(doc, key, number, caption, note):
    path = SCREENSHOTS[key]
    if not path.exists():
        add_para(doc, f"[Evidencia no disponible: {caption}]", italic=True, first_line=False)
        return
    p_num = doc.add_paragraph()
    p_num.paragraph_format.keep_with_next = True
    font(p_num.add_run(f"Figura {number}"), bold=True)
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.keep_with_next = True
    font(p_cap.add_run(caption), italic=True)
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.keep_with_next = True
    shape = pic_p.add_run().add_picture(str(path), width=Inches(6.45))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", f"Figura {number}")
    p_note = doc.add_paragraph()
    font(p_note.add_run("Nota. "), italic=True, size=10)
    font(p_note.add_run(note), size=10)


def make_architecture_diagram():
    path = QA_DIR / "arquitectura_api.png"
    img = Image.new("RGB", (1500, 320), "white")
    draw = ImageDraw.Draw(img)
    try:
        regular = ImageFont.truetype("arial.ttf", 34)
        small = ImageFont.truetype("arial.ttf", 27)
    except OSError:
        regular = ImageFont.load_default()
        small = regular
    labels = ["Postman", "Solicitud HTTP", "Rutas y middleware", "Controlador y validación", "Sanctum", "MySQL"]
    colors = ["#DDF3EC", "#E8EEF7", "#DDF3EC", "#E8EEF7", "#DDF3EC", "#E8EEF7"]
    x, y, w, h, gap = 25, 95, 210, 125, 34
    for i, label in enumerate(labels):
        x0 = x + i * (w + gap)
        draw.rounded_rectangle((x0, y, x0+w, y+h), radius=18, fill=colors[i], outline="#1B5E4B", width=3)
        bbox = draw.multiline_textbbox((0, 0), label, font=small, align="center")
        tw, th = bbox[2]-bbox[0], bbox[3]-bbox[1]
        draw.multiline_text((x0+(w-tw)/2, y+(h-th)/2), label, fill="#17342D", font=small, align="center")
        if i < len(labels)-1:
            ax = x0+w+5
            ay = y+h//2
            draw.line((ax, ay, ax+gap-10, ay), fill="#1B5E4B", width=5)
            draw.polygon([(ax+gap-10, ay), (ax+gap-24, ay-10), (ax+gap-24, ay+10)], fill="#1B5E4B")
    draw.text((25, 22), "Flujo de una solicitud en la API REST de Ecosysgame", fill="#17342D", font=regular)
    img.save(path)
    return path


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)
add_page_number(section.header.paragraphs[0])

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2
normal.paragraph_format.space_after = Pt(0)

for name, size, before, after in [("Heading 1", 12, 12, 0), ("Heading 2", 12, 8, 0), ("Heading 3", 12, 6, 0)]:
    st = styles[name]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.line_spacing = 2
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
styles["Heading 3"].font.italic = True

for style_name in ["List Bullet", "List Number"]:
    st = styles[style_name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.paragraph_format.line_spacing = 2
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(0)

# Portada APA 7 (trabajo estudiantil)
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Validación y pruebas de seguridad de una API REST con Laravel y Postman"), bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Caso de estudio: Ecosysgame"), bold=True)
doc.add_paragraph()
for line in [
    "[Nombre completo del estudiante]",
    "[Nombre de la institución]",
    "CADI Seguridad en Aplicaciones",
    "[Nombre del docente]",
    "3 de septiembre de 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(line))
doc.add_page_break()

add_heading(doc, "Validación y pruebas de seguridad de una API REST con Laravel y Postman", 1)

add_heading(doc, "Introducción", 2)
add_para(doc, "Una interfaz de programación de aplicaciones (API) permite que distintos componentes de software intercambien datos y ejecuten operaciones mediante solicitudes y respuestas. En una API REST, cada recurso se expone a través de endpoints y métodos HTTP. Debido a que estos puntos reciben información y pueden conceder acceso a recursos privados, constituyen una superficie de ataque que debe protegerse mediante validación, autenticación, autorización, limitación de solicitudes y exposición mínima de datos.")
add_para(doc, "El presente informe documenta pruebas controladas realizadas con Postman sobre la API REST de Ecosysgame, desarrollada con Laravel 12. La API utiliza Laravel Sanctum para emitir y validar tokens personales enviados como credenciales Bearer. La documentación oficial indica que estos tokens deben incluirse en el encabezado Authorization y que las rutas protegidas pueden utilizar el middleware auth:sanctum (Laravel, 2026a). Las pruebas se orientaron a verificar el comportamiento funcional y de seguridad ante solicitudes válidas, inválidas, no autenticadas, manipuladas y repetitivas.")

add_heading(doc, "Objetivos", 2)
add_heading(doc, "Objetivo general", 3)
add_para(doc, "Evaluar mediante pruebas controladas con Postman los mecanismos de validación y seguridad implementados en la API REST de Ecosysgame desarrollada con Laravel.")
add_heading(doc, "Objetivos específicos", 3)
for item in [
    "Comprobar el funcionamiento de los endpoints de registro, inicio de sesión, consulta de perfil y cierre de sesión.",
    "Verificar la validación de datos, la política de contraseñas y la unicidad del correo electrónico.",
    "Evaluar la autenticación mediante tokens, la protección de recursos privados y la revocación de credenciales.",
    "Analizar los códigos HTTP y el contenido de las respuestas para identificar controles y posibles debilidades.",
    "Proponer recomendaciones de seguridad sustentadas en los resultados obtenidos.",
]:
    add_bullet(doc, item)

add_heading(doc, "Arquitectura de la API", 2)
add_para(doc, "La solución sigue una arquitectura cliente-servidor. Postman actúa como cliente de pruebas; Laravel recibe la solicitud, aplica las rutas y el middleware, ejecuta el controlador y las reglas de validación, consulta la base de datos y devuelve una respuesta JSON. En los endpoints privados, Sanctum comprueba el token antes de permitir el acceso al controlador.")

diagram = make_architecture_diagram()
p = doc.add_paragraph()
font(p.add_run("Figura 1"), bold=True)
p = doc.add_paragraph()
font(p.add_run("Arquitectura lógica de la API REST evaluada"), italic=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shape = p.add_run().add_picture(str(diagram), width=Inches(6.45))
shape._inline.docPr.set("descr", "Diagrama del flujo Postman, solicitud HTTP, middleware, controlador, Sanctum y MySQL")
shape._inline.docPr.set("title", "Figura 1")
p = doc.add_paragraph()
font(p.add_run("Nota. "), italic=True, size=10)
font(p.add_run("Elaboración propia a partir de la implementación examinada en routes/api.php y AuthController.php."), size=10)

add_heading(doc, "Endpoints evaluados", 2)
add_table(doc,
    ["Método", "Endpoint", "Función", "Protección principal"],
    [
        ("POST", "/api/register", "Registrar una cuenta", "Validación y límite de 10 solicitudes/minuto"),
        ("POST", "/api/login", "Autenticar y emitir token", "Validación, verificación de hash y límite de 5 solicitudes/minuto"),
        ("GET", "/api/profile", "Consultar el perfil", "Token Bearer y middleware auth:sanctum"),
        ("POST", "/api/logout", "Cerrar sesión", "Token Bearer, auth:sanctum y revocación"),
    ], [1100, 1800, 2500, 3960], font_size=9)

add_heading(doc, "Controles de seguridad implementados", 2)
for item in [
    "Validación del lado del servidor para nombre, correo y contraseña.",
    "Correo electrónico único mediante la regla unique:users,email.",
    "Contraseña de mínimo ocho caracteres con letras mayúsculas y minúsculas, números y símbolos.",
    "Almacenamiento de contraseñas mediante hash; la contraseña nunca se devuelve en JSON.",
    "Mensajes genéricos para credenciales incorrectas, con el fin de reducir la enumeración de usuarios.",
    "Autenticación mediante tokens personales de Laravel Sanctum.",
    "Protección de /profile y /logout con auth:sanctum.",
    "Revocación del token usado al cerrar sesión.",
    "Lista explícita de atributos en el registro para impedir que el cliente se asigne privilegios administrativos.",
    "Rate limiting en registro e inicio de sesión.",
    "Configuración APP_DEBUG=false para evitar que una excepción exponga rutas y trazas internas.",
]:
    add_bullet(doc, item)

add_heading(doc, "Procedimiento", 2)
add_heading(doc, "Preparación del entorno", 3)
for step in [
    "Se inició la aplicación Laravel localmente en http://127.0.0.1:8000 y se verificó la conexión con la base de datos MySQL.",
    "Se instaló Laravel Sanctum, se ejecutó la migración de personal_access_tokens y se incorporó HasApiTokens al modelo User.",
    "Se declararon las cuatro rutas de la API y se protegieron profile y logout con auth:sanctum.",
    "Se ejecutaron las pruebas automatizadas del proyecto para comprobar el comportamiento de la API y del portal.",
]:
    add_number(doc, step)

add_heading(doc, "Configuración de Postman", 3)
for step in [
    "Se creó la colección Ecosysgame API - Pruebas de Seguridad.",
    "Se definió la variable de colección base_url con el valor http://127.0.0.1:8000/api.",
    "Se agregó la variable token para almacenar temporalmente el token generado durante el login y revoked_token para comprobar la revocación.",
    "En las solicitudes JSON se configuraron Accept: application/json y Content-Type: application/json. En las rutas privadas se usó Authorization: Bearer {{token}}.",
    "En Scripts > After response se añadieron aserciones para validar código HTTP, estructura JSON, ausencia de campos sensibles y persistencia o eliminación de variables.",
    "Cada solicitud se envió de forma controlada y se registraron el código, el tiempo, el cuerpo de la respuesta y el resultado de las aserciones.",
]:
    add_number(doc, step)
add_para(doc, "Postman permite usar scripts posteriores a la respuesta para ejecutar aserciones y transferir datos entre solicitudes; asimismo, las variables de colección pueden consultarse y modificarse mediante pm.collectionVariables (Postman, 2026a, 2026b).", first_line=True)

add_heading(doc, "Resultados", 2)
add_para(doc, "Se ejecutaron 11 escenarios en Postman. Todos los escenarios realizados produjeron el código esperado y sus aserciones visibles finalizaron correctamente. Además, la suite automatizada del proyecto había registrado 19 pruebas y 86 aserciones aprobadas. La Tabla 2 resume las evidencias disponibles. Los tiempos son los observados en las capturas y pueden variar entre ejecuciones locales.")

rows = [
    ("P01", "Registro válido", "POST /register", "201", "201", "Aprobada", "997 ms"),
    ("P02", "Login correcto", "POST /login", "200", "200", "Aprobada", "947 ms"),
    ("P03", "Perfil con token", "GET /profile", "200", "200", "Aprobada", "335 ms"),
    ("P04", "Perfil sin token", "GET /profile", "401", "401", "Aprobada", "316 ms"),
    ("P05", "Logout correcto", "POST /logout", "200", "200", "Aprobada", "357 ms"),
    ("P06", "Token revocado", "GET /profile", "401", "401", "Aprobada", "418 ms"),
    ("P07", "Correo duplicado", "POST /register", "422", "422", "Aprobada", "469 ms"),
    ("P08", "Contraseña débil", "POST /register", "422", "422", "Aprobada", "485 ms"),
    ("P09", "Credenciales incorrectas", "POST /login", "401", "401", "Aprobada", "969 ms"),
    ("P10", "Asignación de administrador", "POST /register", "201 sin privilegios", "201 sin privilegios", "Aprobada", "872 ms"),
    ("P11", "Demasiados intentos", "POST /login", "429", "429", "Aprobada", "282 ms"),
]
add_table(doc, ["ID", "Prueba", "Solicitud", "Esperado", "Obtenido", "Estado", "Tiempo"], rows,
          [600, 2100, 1700, 1250, 1250, 1300, 1160], font_size=8)

add_heading(doc, "Evidencias representativas", 3)
add_figure(doc, "register", 2, "Registro válido con respuesta 201 Created", "La respuesta confirma la creación de la cuenta, omite la contraseña y muestra tres aserciones aprobadas. Los datos pertenecen exclusivamente al entorno de pruebas.")
add_figure(doc, "login", 3, "Inicio de sesión correcto y emisión de token", "Se obtuvo 200 OK y cuatro aserciones aprobadas. El token mostrado era temporal y fue revocado posteriormente durante la prueba de logout.")
add_figure(doc, "profile", 4, "Consulta autorizada del perfil", "El endpoint devolvió únicamente id, name, email y created_at; no expuso password, remember_token ni is_admin.")
add_figure(doc, "no_token", 5, "Rechazo del perfil sin token", "La API respondió 401 Unauthorized y no entregó información del usuario.")
add_figure(doc, "weak", 6, "Rechazo de contraseña débil", "La API respondió 422 y señaló incumplimientos de complejidad; las tres aserciones finalizaron correctamente.")
add_figure(doc, "mass_assignment", 7, "Intento de asignación de privilegios administrativos", "Aunque la cuenta fue creada, los campos is_admin, role y permissions enviados por el cliente fueron ignorados y no aparecieron en la respuesta.")
add_figure(doc, "logout", 8, "Cierre de sesión y revocación del token", "La API confirmó el cierre de sesión con 200 OK y las cuatro aserciones aprobaron la conservación temporal y posterior eliminación de la variable activa.")
add_figure(doc, "rate_limit", 9, "Bloqueo por exceso de intentos", "Después de solicitudes repetitivas, la API respondió 429 Too Many Requests, incluyó Retry-After y devolvió un mensaje mínimo sin traza interna gracias a APP_DEBUG=false.")

add_heading(doc, "Cobertura pendiente frente a la guía", 3)
add_para(doc, "La guía suministrada propone 14 pruebas base. Las evidencias recopiladas cubren los controles centrales, pero cuatro variantes no fueron documentadas como solicitudes independientes. Por rigor metodológico, se presentan como pendientes y no se les atribuye un resultado no observado.")
pending = [
    ("T02", "Registro con formato de correo inválido", "422", "Pendiente de captura independiente"),
    ("T03", "Registro con campos obligatorios ausentes", "422", "Pendiente de captura independiente"),
    ("T08", "Login con correo de usuario inexistente", "401 y mensaje genérico", "Pendiente de comparación con contraseña incorrecta"),
    ("T10", "Perfil con token arbitrario, por ejemplo 123456789", "401", "Pendiente de captura independiente"),
]
add_table(doc, ["Guía", "Caso", "Resultado esperado", "Estado documental"], pending,
          [900, 3600, 2200, 2660], font_size=9)

add_heading(doc, "Análisis de los resultados", 2)
add_heading(doc, "Validación de entrada", 3)
add_para(doc, "Los resultados 422 obtenidos ante el correo duplicado y la contraseña débil demuestran que la validación se ejecuta en el servidor. Esto es esencial porque un cliente puede omitir el formulario web y enviar una solicitud directa. La implementación utiliza reglas explícitas para el nombre, el formato y unicidad del correo, la confirmación de contraseña y su complejidad. El registro crea el usuario únicamente con los campos validados, lo que reduce el riesgo de asignación masiva.")
add_heading(doc, "Autenticación y autorización", 3)
add_para(doc, "El login correcto entregó un token Bearer y el perfil autorizado respondió 200. En contraste, el mismo endpoint rechazó con 401 las solicitudes sin token y con un token revocado. Esto confirma que la identidad se comprueba antes de acceder al recurso. La autenticación determina quién realiza la solicitud; la autorización establece si esa identidad puede ejecutar una acción. En esta API, auth:sanctum aporta la barrera de autenticación, mientras que la lista de campos de salida limita la información entregada.")
add_heading(doc, "Protección de datos", 3)
add_para(doc, "Las respuestas de registro, login y perfil no contienen password, remember_token ni is_admin. La contraseña se almacena mediante hash y nunca se devuelve al cliente. La minimización de datos disminuye el impacto de una respuesta interceptada o registrada accidentalmente. La prueba de manipulación también mostró que enviar is_admin=true no concede privilegios.")
add_heading(doc, "Revocación y limitación de intentos", 3)
add_para(doc, "Después del logout, el token previamente válido dejó de autorizar el perfil, lo que demuestra una revocación efectiva. Asimismo, el sexto intento de login dentro de la ventana configurada produjo 429 Too Many Requests y el encabezado Retry-After. Este control dificulta ataques automatizados de fuerza bruta. OWASP incluye los fallos de autenticación y el consumo no restringido de recursos entre los riesgos relevantes para APIs (OWASP Foundation, 2023).")
add_heading(doc, "Configuración segura de errores", 3)
add_para(doc, "Durante la preparación se observó que el modo de depuración podía incluir una traza extensa en la respuesta 429. Antes de conservar la evidencia final se configuró APP_DEBUG=false y se limpió la caché de configuración. La respuesta definitiva quedó reducida al mensaje Too Many Attempts. La captura insegura fue excluida del informe. Este ajuste evita revelar rutas locales, nombres de clases y detalles internos útiles para un atacante.")

add_heading(doc, "Matriz de controles", 2)
controls = [
    ("Validación de email", "Parcialmente evidenciado", "Existe en código; falta captura específica de formato inválido"),
    ("Política de contraseña", "Sí", "P08: 422 ante contraseña débil"),
    ("Email único", "Sí", "P07: 422 ante correo duplicado"),
    ("Autenticación", "Sí", "P02: token Bearer; P09: 401 ante credenciales incorrectas"),
    ("Protección del perfil", "Sí", "P03: 200 con token; P04: 401 sin token"),
    ("Revocación de token", "Sí", "P05 y P06"),
    ("Rate limiting", "Sí", "P11: 429 y Retry-After"),
    ("Protección de datos sensibles", "Sí", "P03: campos sensibles ausentes"),
    ("Prevención de escalamiento", "Sí", "P10: campos administrativos ignorados"),
    ("Errores sin trazas internas", "Sí", "P11 con APP_DEBUG=false"),
]
add_table(doc, ["Control", "Estado", "Evidencia"], controls, [2600, 2200, 4560], font_size=9)

add_heading(doc, "Hallazgos", 2)
findings = [
    ("H01", "El modo debug expuso temporalmente una traza durante la prueba de límite.", "Configuración de depuración habilitada en entorno local.", "Divulgación de rutas y componentes internos.", "Media", "Configurar APP_DEBUG=false fuera de desarrollo y verificar respuestas de excepción.", "Corregido"),
    ("H02", "La guía exige cuatro variantes aún sin captura independiente.", "Cobertura manual incompleta.", "Evidencia insuficiente para afirmar cobertura total del procedimiento.", "Baja", "Ejecutar T02, T03, T08 y T10 y anexar sus capturas.", "Pendiente"),
    ("H03", "Los tokens personales no muestran una política de expiración en las evidencias.", "Sanctum admite tokens persistentes hasta su revocación, según configuración.", "Un token filtrado podría conservar validez durante más tiempo del necesario.", "Media", "Definir expiración, rotación y capacidades mínimas para tokens cuando la API evolucione.", "Recomendación"),
]
add_table(doc, ["ID", "Hallazgo", "Causa", "Impacto", "Severidad", "Recomendación", "Estado"], findings,
          [500, 1800, 1500, 1600, 900, 2160, 900], font_size=7.5)

add_heading(doc, "Reflexión", 2)
questions = [
    ("¿Por qué una API representa una superficie de ataque?", "Porque expone operaciones y datos a clientes que pueden enviar solicitudes directas, manipuladas o automatizadas, sin depender de la interfaz web."),
    ("¿Por qué no basta validar en el frontend?", "El frontend está bajo control del cliente y puede modificarse o evitarse. La API debe validar cada solicitud en el servidor."),
    ("¿Qué diferencia existe entre autenticación y autorización?", "La autenticación comprueba la identidad; la autorización decide qué recursos o acciones se permiten a esa identidad."),
    ("¿Por qué un endpoint privado debe rechazar solicitudes sin token?", "Porque, sin una credencial válida, no existe evidencia de identidad ni una base para conceder acceso al recurso."),
    ("¿Por qué no se deben devolver contraseñas?", "La contraseña es un secreto de autenticación. Exponerla, incluso cifrada o resumida, incrementa el riesgo de compromiso y ataques fuera de línea."),
    ("¿Qué ocurre si el login admite intentos ilimitados?", "Facilita ataques de fuerza bruta y relleno de credenciales, además de consumir recursos del servidor."),
    ("¿Por qué son importantes los códigos HTTP?", "Permiten interpretar de forma uniforme si una operación tuvo éxito, falló por validación, carece de autenticación o fue limitada."),
    ("¿Qué responsabilidad ética tiene el desarrollador?", "Debe minimizar los datos recolectados y expuestos, aplicar controles razonables, corregir hallazgos y evitar prácticas que pongan en riesgo a los usuarios."),
]
for q, a in questions:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    font(p.add_run(q + " "), bold=True)
    font(p.add_run(a))

add_heading(doc, "Propuestas de mejora", 2)
improvements = [
    ("Tokens con vigencia indefinida", "Expiración, rotación y capacidades mínimas", "Configurar la expiración de Sanctum y emitir tokens con abilities acordes con cada cliente.", "Reduce la ventana de uso de un token filtrado y limita su alcance."),
    ("Pruebas manuales no completamente repetibles", "Automatización de la colección", "Completar los cuatro casos pendientes, usar datos dinámicos y ejecutar la colección en CI con Postman CLI o Newman.", "Evidencia reproducible y detección temprana de regresiones."),
    ("Riesgo de configuraciones inseguras", "Verificación de configuración por ambiente", "Mantener APP_DEBUG=false en producción, proteger secretos en .env y añadir comprobaciones al despliegue.", "Evita trazas, credenciales y detalles internos en respuestas públicas."),
]
add_table(doc, ["Problema", "Control propuesto", "Implementación en Laravel/Postman", "Resultado esperado"], improvements,
          [2100, 2200, 2800, 2260], font_size=8.5)

add_heading(doc, "Conclusiones", 2)
add_para(doc, "Las pruebas ejecutadas demuestran que la API de Ecosysgame aplica controles efectivos de validación, autenticación mediante Sanctum, protección de endpoints, revocación de tokens, minimización de datos, prevención de asignación de privilegios y limitación de intentos. Los códigos 200, 201, 401, 422 y 429 coincidieron con los comportamientos esperados en los 11 escenarios documentados.")
add_para(doc, "No obstante, afirmar que una API es segura requiere más que comprobar rutas exitosas. Las solicitudes negativas y manipuladas permitieron validar que los controles se mantienen cuando el cliente se comporta de forma incorrecta o maliciosa. La corrección del modo de depuración mostró además que la configuración del ambiente forma parte de la seguridad y no debe separarse del código.")
add_para(doc, "La evidencia disponible permite concluir que la implementación presenta una base adecuada para el laboratorio. Sin embargo, el cumplimiento documental total de la guía queda condicionado a ejecutar y capturar cuatro variantes adicionales: correo inválido, campos faltantes, usuario inexistente y token arbitrario inválido. Después de incorporar esas evidencias y repetir la colección de manera controlada, el informe podrá considerarse completo frente al procedimiento solicitado.")

add_heading(doc, "Referencias", 2)
refs = [
    "Guía de pruebas de seguridad en API REST (Laravel y Postman). (s. f.). [Material de clase].",
    "Laravel. (2026a). Laravel Sanctum. https://laravel.com/docs/12.x/sanctum",
    "Laravel. (2026b). Validation. https://laravel.com/docs/12.x/validation",
    "OWASP Foundation. (2023). OWASP API Security Top 10 – 2023. https://owasp.org/API-Security/editions/2023/en/0x11-t10/",
    "Postman. (2026a). Use scripts to add logic and tests to Postman requests. https://learning.postman.com/docs/tests-and-scripts/write-scripts/intro-to-scripts/",
    "Postman. (2026b). Reference variables in Postman scripts. https://learning.postman.com/docs/tests-and-scripts/write-scripts/postman-sandbox-reference/pm-variables/",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    font(p.add_run(ref))

doc.core_properties.title = "Validación y pruebas de seguridad de una API REST con Laravel y Postman"
doc.core_properties.subject = "Informe técnico de seguridad - Ecosysgame"
doc.core_properties.author = "[Nombre completo del estudiante]"
doc.core_properties.keywords = "Laravel, Postman, API REST, Sanctum, seguridad, Ecosysgame"
doc.save(OUT)
print(OUT)
